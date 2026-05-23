#!/usr/bin/env python3
"""Create a Word draft from the paper Markdown with PNG figures embedded.

The output is still a transfer draft, not the official journal template. It is
intended to reduce manual Word/WPS work by preserving headings, tables, formulas
as text blocks, and generated figures. With ``--journal-style`` it also adds
front-matter placeholders, equation numbers, and three-line table borders.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
CAPTION_RE = re.compile(
    r"^(图|表)\s*\d+\s+(?!(为|给出|显示|进一步|所示|见)\b)"
    r"|^(Fig\.|Table)\s*\d+\b"
)
ORDERED_LIST_RE = re.compile(r"^\d+\.\s+")


def clean_inline(text: str) -> str:
    text = LINK_RE.sub(r"\1", text)
    text = re.sub(r"\$([^$]+)\$", lambda m: clean_formula(m.group(1)), text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def set_run_font(run, size: int = 10, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(  # noqa: SLF001
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "宋体",
    )
    run.font.size = Pt(size)
    run.bold = bold


ENGLISH_TITLE = (
    "Named Entity Recognition Method for Red Jujube Cultivation Based on "
    "Expert Dictionary and Boundary Prediction"
)


def add_paragraph(doc: Document, text: str, size: int = 10, bold: bool = False, align=None):
    if not text:
        return None
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(clean_inline(text))
    set_run_font(run, size=size, bold=bold)
    return p


def is_caption(text: str) -> bool:
    return bool(CAPTION_RE.match(text.strip()))


def add_caption(doc: Document, text: str) -> None:
    add_paragraph(doc, text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def clean_formula(text: str) -> str:
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\\\\s*$", "", text)
    replacements = [
        (r"\operatorname{MacBERT}", "MacBERT"),
        (r"\mathbb{R}", "R"),
        (r"\in", "∈"),
        (r"\times", "×"),
        (r"\ldots", "…"),
        (r"\{", "{"),
        (r"\}", "}"),
        (r"\gamma", "γ"),
        (r"\alpha_t", "α_t"),
        (r"\log", "log"),
        (r"\begin{cases}", "{"),
        (r"\end{cases}", ""),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    text = re.sub(r"\\[;,!]", " ", text)
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    text = text.replace("  ", " ")
    text = text.replace(" & ", "    ")
    text = text.replace("&", "    ")
    return text.strip()


def add_formula(doc: Document, eq_lines: list[str], eq_no: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)

    formula_lines = [clean_formula(line) for line in eq_lines if line.strip()]
    formula_text = "\n".join(formula_lines)
    run = p.add_run()
    set_run_font(run, size=10)
    for idx, part in enumerate(formula_text.splitlines()):
        if idx:
            run.add_break(WD_BREAK.LINE)
        run.add_text(part)

    num_run = p.add_run(f"\t({eq_no})")
    set_run_font(num_run, size=10)


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 0:
        add_paragraph(doc, text, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    add_paragraph(doc, text, size=12 if level == 1 else 11, bold=True)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def apply_three_line_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = list(table.rows)
    if not rows:
        return
    nil = {"val": "nil"}
    solid = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for row in rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=nil,
                left=nil,
                bottom=nil,
                right=nil,
                insideH=nil,
                insideV=nil,
            )
    for cell in rows[0].cells:
        set_cell_border(cell, top=solid, bottom=solid, left=nil, right=nil)
    for cell in rows[-1].cells:
        set_cell_border(cell, bottom=solid, left=nil, right=nil)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
        is_sep = all(re.fullmatch(r":?-{2,}:?", p or "") for p in parts)
        if not is_sep:
            rows.append([clean_inline(p) for p in parts])
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]], journal_style: bool = False) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9, bold=(r_idx == 0))
    if journal_style:
        apply_three_line_table(table)


def add_front_matter_placeholders(doc: Document) -> None:
    add_paragraph(
        doc,
        ENGLISH_TITLE,
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(doc, "作者：待作者补充", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "作者单位：待作者补充二级单位、城市和邮编",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(doc, "Authors: To be completed", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "Affiliations: To be completed",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "中图分类号：待作者确认    文献标识码：A    文章编号：编辑部填写",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "基金项目：待作者补充    作者简介/通信作者：待作者补充",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def resolve_image(md_path: Path, image_ref: str) -> Path:
    raw = (md_path.parent / image_ref).resolve()
    if raw.suffix.lower() == ".svg":
        png = md_path.parent / "figures_png" / (raw.stem + ".png")
        if png.exists():
            return png.resolve()
    return raw


def build_docx(md_path: Path, out_path: Path, journal_style: bool = False) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    buf: list[str] = []
    i = 0
    eq_no = 1
    title_seen = False

    def flush() -> None:
        nonlocal buf
        if buf:
            add_paragraph(doc, " ".join(x.strip() for x in buf), size=10)
            buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush()
            i += 1
            continue
        if stripped.startswith("|"):
            flush()
            rows, i = parse_table(lines, i)
            add_table(doc, rows, journal_style=journal_style)
            continue
        if is_caption(stripped):
            flush()
            add_caption(doc, stripped)
            i += 1
            continue
        if ORDERED_LIST_RE.match(stripped):
            flush()
            add_paragraph(doc, stripped, size=10)
            i += 1
            continue
        if stripped.startswith("# "):
            flush()
            add_heading(doc, stripped[2:], 0)
            if journal_style and not title_seen:
                add_front_matter_placeholders(doc)
                title_seen = True
            i += 1
            continue
        if stripped.startswith("## "):
            flush()
            add_heading(doc, stripped[3:], 1)
            i += 1
            continue
        if stripped.startswith("### "):
            flush()
            add_heading(doc, stripped[4:], 2)
            i += 1
            continue
        if stripped.startswith("#### "):
            flush()
            add_heading(doc, stripped[5:], 3)
            i += 1
            continue
        if stripped.startswith("$$"):
            flush()
            eq_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("$$"):
                eq_lines.append(lines[i].strip())
                i += 1
            if journal_style:
                add_formula(doc, eq_lines, eq_no)
                eq_no += 1
            else:
                eq_text = " ".join(eq_lines)
                eq_text = "[公式] " + eq_text
                add_paragraph(doc, eq_text, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            if i < len(lines):
                i += 1
            continue
        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush()
            alt, ref = image_match.groups()
            image_path = resolve_image(md_path, ref)
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(image_path), width=Cm(15))
            else:
                add_paragraph(doc, f"[图件缺失：{alt}；文件：{ref}]", size=10)
            i += 1
            continue
        buf.append(line)
        i += 1
    flush()
    doc.save(out_path)


def main() -> int:
    args = [arg for arg in sys.argv[1:] if arg != "--journal-style"]
    journal_style = "--journal-style" in sys.argv[1:]
    if len(args) != 2:
        print(
            "Usage: md_to_docx_with_images.py [--journal-style] input.md output.docx",
            file=sys.stderr,
        )
        return 2
    build_docx(Path(args[0]), Path(args[1]), journal_style=journal_style)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
